# import numpy as np
import pandas as pd
from docx import Document
import datetime
import os
import sys
import client_details

def resource_path(relative_path):
    """Get abolute path to resource, works for dev and for Pyinstaller"""
    try:
        # Pyinstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path,relative_path)

def generate_invoices(invoice_month_num,invoice_year,filename,sheet_name):#, template_file):
    invoice_month_num = int(invoice_month_num)
    invoice_year = int(invoice_year)
    month_name = datetime.date(1900, invoice_month_num, 1).strftime('%B')

    # Template
    template_file = resource_path('Invoice Template.docx')

    # Import data
    df = pd.read_excel(filename, index_col=None, engine="openpyxl", skiprows=1, sheet_name=sheet_name)  
    
    # Only Read useful Columns
    df = df[['Day','Date','Place worked Trim','Start','Finish','Hourly Rate','Hours worked','Daily Salary','Monthly Salary']]
    
    # Create Month number/Year/Date field
    df['month_num'] = pd.to_datetime(df['Date']).dt.month
    df['year'] = pd.to_datetime(df['Date']).dt.year
    df['Date'] = df['Date'].dt.date
    
    # Create Shift Pattern field
    df['Shift_Pattern'] = df['Start'].apply(str).apply(lambda x: x[0:5]) + " - " + df['Finish'].apply(str).apply(lambda x: x[0:5])

    # Filter on Month Number and Year
    df = df[df['month_num'] == invoice_month_num]
    df = df[df['year'] == invoice_year]

    # Store All Clients
    unique_clients = df['Place worked Trim'].dropna().unique()

    # Main loop and invoice generation
    for client in unique_clients:
        
        #open pdf or word file template
        doc = Document(template_file)
        doc.tables                                              #a list of all tables in document
        invoice_table = doc.tables[2]

        # Get specific pharmacy data
        df_pharm = df[df['Place worked Trim'] == client]
        
        # Make date display saturday if in ellesmere
        def saturday(day,date):
            if day == 'Saturday': return str(date) + ' (Saturday)'
            else: return date
        if client == "Ellesmere":
            df_pharm.loc[:,'Date'] = df_pharm.apply(lambda row: saturday(row['Day'],row['Date']),axis=1 )
            # df_pharm['Date'] = df_pharm.apply(lambda row: saturday(row['Day'],row['Date']),axis=1 )            
        
        # select only data we want to display 
        df_pharm = df_pharm[['Date','Hours worked','Shift_Pattern','Hourly Rate','Daily Salary']]
        
        #add invoice total on sum at end of dataframe
        pharm_sum = df_pharm['Daily Salary'].sum()
        sum_data = {'Date': ['Total'], 'Hours worked':[''], 'Shift_Pattern':[''], 'Hourly Rate':[''],'Daily Salary':[pharm_sum]}
        sum_df = pd.DataFrame(data=sum_data)
        df_pharm = df_pharm.append(sum_df,ignore_index=True)
        
        # clear table
        for row in range(len(doc.tables[2].rows)-1):  #rows
            for col in range(5):  #columns
                invoice_table.cell(row+1,col).text = ''
        
        # write data to table
        for row in range(len(df_pharm.index)):  #rows 
            for col in range(5):  #columns
                while True:
                    try:
                        invoice_table.cell(row+1,col).text = str(df_pharm.values[row,col])
                        if row == len(df_pharm.index) - 1:  # bold last row
                            invoice_table.cell(row+1,col).paragraphs[0].runs[0].font.bold = True
                        break
                    except IndexError: # add row if not enough rows
                        print("IndexError raised, adding new row to table")
                        new_row = invoice_table.add_row().cells
                        # invoice_table.style = 'Table Grid'

        # Format Other parts of invoice document
        todays_date = datetime.datetime.today().strftime('%Y-%m-%d')
        doc.tables[0].cell(0,1).text = todays_date          # Date a begining of document
        doc.paragraphs[7].add_run(todays_date)              # Date at end of document
        # Pharmacy Branch
        if client == client_details.client_name_1:    doc.tables[1].cell(0,1).text = client_details.branch_name_1
        elif client == client_details.client_name_2:   doc.tables[1].cell(0,1).text = client_details.branch_name_2
        else:                       doc.tables[1].cell(0,1).text = client                # Pharmacy Branch

        
        doc_name = str(invoice_year) + " " + month_name + " - " + client + ' Invoice'  + '.docx' 
        doc_name = doc_name.replace("/", "-").replace("\\", "-").replace(":", "-").replace("?", "-").replace('"', "-").replace('"', "-").replace("<", "-").replace(">", "-").replace("|", "-").replace("*", "-")
        doc.save(doc_name)
